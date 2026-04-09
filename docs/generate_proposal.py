"""Generate NSPD Ghana v2 Proposal as a Word document (non-technical, client-facing)."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page Setup ──────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    h.font.bold = True
    if level == 1:
        h.font.size = Pt(20)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(15)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.font.size = Pt(12)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

BLUE = RGBColor(0x00, 0x33, 0x66)
GOLD = RGBColor(0xCC, 0xA3, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_BG = RGBColor(0xF4, 0xF7, 0xF6)


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(sh)


def add_styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, "003366")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F4F7F6")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph("")  # spacing
    return table


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


# ══════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("NSPD GHANA")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = BLUE
run.font.name = "Calibri"

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("System Enhancement Proposal")
run.font.size = Pt(22)
run.font.color.rgb = GRAY
run.font.name = "Calibri"

doc.add_paragraph("")

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run("Version 2.0")
run.font.size = Pt(14)
run.font.color.rgb = GOLD
run.bold = True

doc.add_paragraph("")
doc.add_paragraph("")

details = [
    ("Prepared for:", "Ghana Maritime Authority"),
    ("Prepared by:", "NSPD Ghana Development Team"),
    ("Date:", "April 2026"),
    ("Document Version:", "1.0"),
    ("Classification:", "Confidential"),
]
for label, value in details:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + " ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    run = p.add_run(value)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════

doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1.  Executive Summary",
    "2.  What the System Does Today",
    "3.  What Needs to Change",
    "4.  Proposed New Features",
    "     4.1  Phase 1 — Core Capabilities",
    "     4.2  Phase 2 — Compliance & Records",
    "     4.3  Phase 3 — Operational Tools",
    "     4.4  Phase 4 — Finishing Touches",
    "5.  Implementation Timeline",
    "6.  Summary of Deliverables",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════

doc.add_heading("1. Executive Summary", level=1)

doc.add_paragraph(
    "The National Seafarer Placement Database (NSPD) Ghana system currently serves as an "
    "internal dashboard for reviewing seafarer training applications. Staff can log in to view "
    "submissions, browse analytics, and export data — but the system is read-only. There is no way "
    "for seafarers to apply through the platform, no process for reviewing and approving applications, "
    "and no tools for managing user accounts or tracking documents."
)

doc.add_paragraph(
    "This proposal outlines a plan to transform NSPD Ghana into a complete, end-to-end seafarer "
    "registration and compliance management platform. The upgrade is organised into four phases, "
    "each delivering working features that can be used immediately upon completion."
)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("In summary, Version 2 will enable the system to:")
run.bold = True

benefits = [
    "Accept new seafarer applications directly through the platform",
    "Process applications through a structured review and approval workflow",
    "Store and manage certificates, medical records, and identification documents",
    "Track certification expiry dates and flag compliance issues automatically",
    "Provide a full audit trail of all system activity",
    "Send email notifications to seafarers and staff at key milestones",
    "Support user account management without requiring developer involvement",
]
for b in benefits:
    add_bullet(b)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  2. WHAT THE SYSTEM DOES TODAY
# ══════════════════════════════════════════════════════════

doc.add_heading("2. What the System Does Today", level=1)

doc.add_paragraph(
    "Before outlining what needs to change, it is important to acknowledge what the current system "
    "already does well. The existing platform provides a solid foundation to build upon."
)

add_styled_table(
    ["Feature", "Description"],
    [
        ["Secure Login", "Staff log in with a username and password. Sessions are protected against common web security threats."],
        ["Dashboard", "A summary view showing total submissions, average sea experience, the most common rank, and recent activity."],
        ["Submissions List", "A searchable, paginated list of all seafarer applications, with filtering by rank."],
        ["Submission Details", "A full-page view of any individual application, showing personal details, qualifications, and sea experience."],
        ["Reports & Charts", "Five interactive charts covering registration trends, rank distribution, experience levels, certification coverage, and medical fitness."],
        ["PDF Export", "Download an individual submission as a formatted PDF report."],
        ["CSV Export", "Export all submissions (or a filtered subset) as a spreadsheet file."],
        ["User Roles", "Three access levels — Administrator, Reviewer, and Viewer — controlling who can export data."],
    ],
    col_widths=[1.8, 4.7],
)

doc.add_paragraph(
    "The system currently holds 574 seafarer application records that were imported from a previous data source."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  3. WHAT NEEDS TO CHANGE
# ══════════════════════════════════════════════════════════

doc.add_heading("3. What Needs to Change", level=1)

doc.add_paragraph(
    "A thorough review of the platform has identified six significant gaps that limit its usefulness "
    "as an operational tool for the Ghana Maritime Authority."
)

gaps = [
    (
        "No way for seafarers to apply. ",
        "The system is staff-only. There is no online form for seafarers to submit applications. "
        "All existing records were loaded manually from an external source."
    ),
    (
        "No application review process. ",
        "Once an application is in the system, there is no way to mark it as reviewed, approved, or rejected. "
        "There is no assignment of applications to specific reviewers, and no way to leave internal notes."
    ),
    (
        "No user management. ",
        "New staff accounts can only be created by a developer with direct access to the database. "
        "Administrators cannot add, edit, or remove users through the platform."
    ),
    (
        "No document storage. ",
        "The system records whether a seafarer has certifications or medical clearance, but only as a "
        "simple Yes or No. Actual documents (certificates, medical reports, ID scans) are not captured or stored."
    ),
    (
        "Limited seafarer information. ",
        "Important details are missing from the data — date of birth, nationality, identification numbers, "
        "emergency contacts, and residential address. Certifications and medical records lack dates, "
        "so the system cannot flag expired documents."
    ),
    (
        "No activity tracking. ",
        "There is no record of who accessed what, when exports were performed, or who made decisions on applications. "
        "This makes it difficult to maintain accountability."
    ),
]

for bold_part, rest in gaps:
    add_bullet(rest, bold_prefix=bold_part)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  4. PROPOSED NEW FEATURES
# ══════════════════════════════════════════════════════════

doc.add_heading("4. Proposed New Features", level=1)

doc.add_paragraph(
    "The proposed enhancements are organised into four phases. Each phase delivers a complete set "
    "of features that can be deployed and used as soon as it is finished — there is no need to "
    "wait for the entire project to be completed before seeing results."
)

# ── Phase 1 ─────────────────────────────────────────────

doc.add_heading("Phase 1 — Core Capabilities", level=2)

doc.add_paragraph(
    "These are the most important additions. They address the fundamental gaps and transform the "
    "system from a passive data viewer into an active operational tool."
)

doc.add_heading("Online Seafarer Registration Form", level=3)

doc.add_paragraph(
    "A public-facing, multi-step application form that seafarers can access without needing a login. "
    "The form will guide applicants through a clear, step-by-step process:"
)

steps = [
    ("Personal Information — ", "Name, date of birth, contact details, nationality, and identification number."),
    ("Qualifications — ", "Maritime rank, training courses completed, and ISPS/GMA familiarisation status."),
    ("Sea Experience — ", "Voyage history, total years at sea, and most recent vessel type."),
    ("Medical Fitness — ", "Current medical status, examining clinic, and certificate expiry date."),
    ("Document Uploads — ", "Certificates, medical reports, identification scans, and a passport-size photo."),
    ("Review & Submit — ", "A summary of all entered information with the ability to go back and make corrections before final submission."),
]
for bold_part, rest in steps:
    add_bullet(rest, bold_prefix=bold_part)

doc.add_paragraph(
    "Upon successful submission, the seafarer receives a confirmation page with a unique reference "
    "number. The system automatically checks for duplicate applications using the email address "
    "and phone number."
)

doc.add_heading("Application Review & Approval Workflow", level=3)

doc.add_paragraph(
    "A structured process for reviewing and making decisions on applications. "
    "Each application will move through clear stages:"
)

add_styled_table(
    ["Stage", "What Happens"],
    [
        ["Pending", "The application has been submitted and is waiting to be picked up by a reviewer."],
        ["Under Review", "A reviewer has been assigned and is examining the application and supporting documents."],
        ["Approved", "The application has been reviewed and accepted."],
        ["Rejected", "The application has been reviewed and declined, with a reason provided."],
    ],
    col_widths=[1.5, 5.0],
)

doc.add_paragraph(
    "Reviewers will be able to leave internal notes on any application. When rejecting an application, "
    "a reason must be provided. The dashboard will be updated to show how many applications are "
    "awaiting review and how many have been processed today."
)

doc.add_heading("User Management", level=3)

doc.add_paragraph(
    "A dedicated page for administrators to manage staff accounts. This includes the ability to "
    "create new users, assign roles (Administrator, Reviewer, or Viewer), reset passwords, and "
    "deactivate accounts that are no longer needed. Staff members will also be able to change "
    "their own password from within the platform."
)

doc.add_heading("Document Upload & Storage", level=3)

doc.add_paragraph(
    "The system will accept and securely store documents uploaded during the application process "
    "or added later by staff. Supported documents include:"
)

doc_types = [
    "Certificates of Competency and Proficiency",
    "STCW training certificates",
    "Medical fitness certificates",
    "Passport or Ghana Card scans",
    "Passport-size photographs",
    "Other supporting documents",
]
for d in doc_types:
    add_bullet(d)

doc.add_paragraph(
    "Documents will be viewable and downloadable from the submission detail page. The submissions "
    "list will show at a glance whether required documents have been uploaded or are still missing."
)

doc.add_page_break()

# ── Phase 2 ─────────────────────────────────────────────

doc.add_heading("Phase 2 — Compliance & Records Management", level=2)

doc.add_paragraph(
    "These features expand the information the system captures, enabling real compliance monitoring "
    "rather than simple Yes/No tracking."
)

doc.add_heading("Enhanced Seafarer Profiles", level=3)

doc.add_paragraph(
    "The seafarer record will be expanded to include date of birth, gender, nationality, "
    "identification number (Ghana Card or Passport), residential address, region, and emergency "
    "contact information. A passport-size photo will be stored alongside each profile. "
    "This brings the system in line with the data requirements of the IMO and the Ghana Maritime Authority."
)

doc.add_heading("Certification Tracking", level=3)

doc.add_paragraph(
    "Instead of recording certifications as a simple Yes or No, the system will track each "
    "certificate individually — including the certificate name, number, issuing authority, "
    "date of issue, and date of expiry. The uploaded certificate document will be linked directly "
    "to each record."
)

doc.add_paragraph(
    "The system will automatically monitor expiry dates and flag certificates that are expiring "
    "within 30, 60, or 90 days. A new alert section on the dashboard will highlight seafarers "
    "whose certifications require renewal, helping the Authority maintain fleet-wide compliance."
)

doc.add_heading("Voyage & Employment History", level=3)

doc.add_paragraph(
    "A detailed record of each seafarer's voyages — including vessel name, ship type, shipping "
    "company, rank held, and dates of service. The system will automatically calculate total sea "
    "experience from these records, replacing the current manual entry. This provides a verifiable "
    "sea service history that can support Certificate of Competency applications."
)

doc.add_heading("Medical Records", level=3)

doc.add_paragraph(
    "Medical fitness will be tracked with full detail — examination date, expiry date, examining "
    "doctor and clinic, result (Fit, Unfit, or Fit with Restrictions), and the uploaded medical "
    "certificate. The system will alert staff when a seafarer's medical clearance is nearing "
    "expiry, ensuring no one is deployed without valid medical fitness."
)

doc.add_page_break()

# ── Phase 3 ─────────────────────────────────────────────

doc.add_heading("Phase 3 — Operational Tools", level=2)

doc.add_paragraph(
    "These features improve day-to-day efficiency for staff using the system."
)

doc.add_heading("Advanced Search & Filtering", level=3)

doc.add_paragraph(
    "A more powerful search system that allows staff to find records by any combination of criteria — "
    "name, rank, application status, date range, ship type, experience level, medical status, and "
    "certification status. Filters can be saved as presets for frequently used searches "
    "(e.g., \"Pending Engine Cadets\" or \"Expiring Medicals\"). Table columns will be sortable "
    "with a single click, and the number of results shown per page will be configurable."
)

doc.add_heading("Activity & Audit Log", level=3)

doc.add_paragraph(
    "A complete record of all significant actions taken within the system — who logged in, who "
    "viewed or exported which records, and who approved or rejected which applications. "
    "This log will be accessible to administrators and can be filtered by staff member, action "
    "type, and date range. It provides the accountability and traceability expected of a "
    "government information system."
)

doc.add_heading("Email Notifications", level=3)

doc.add_paragraph(
    "Automated email notifications at key points in the process:"
)

notifications = [
    ("To seafarers — ", "Confirmation when their application is submitted, and notification when their application is approved or rejected."),
    ("To reviewers — ", "Alert when a new application is assigned to them for review."),
    ("To administrators — ", "A weekly summary of pending applications and overall system activity."),
]
for bold_part, rest in notifications:
    add_bullet(rest, bold_prefix=bold_part)

doc.add_heading("Administration Settings", level=3)

doc.add_paragraph(
    "A settings page for administrators to control system behaviour — such as opening or closing "
    "the application form, configuring email delivery, and managing dropdown options (rank lists, "
    "ship types) without needing to involve a developer."
)

doc.add_page_break()

# ── Phase 4 ─────────────────────────────────────────────

doc.add_heading("Phase 4 — Finishing Touches", level=2)

doc.add_paragraph(
    "These features add polish and convenience to the platform."
)

doc.add_heading("Enhanced Reporting & Printing", level=3)

doc.add_paragraph(
    "Print-ready views for submission detail pages. The ability to generate PDF reports for "
    "multiple submissions at once (batch export). A seafarer ID card template with photo, name, "
    "rank, registration number, and a QR code. Charts on the reports page will include a "
    "\"Download as Image\" option for use in presentations and external reports."
)

doc.add_heading("Data Import Tool", level=3)

doc.add_paragraph(
    "A tool for importing seafarer records from spreadsheet files (CSV). Staff will be able to "
    "upload a file, map the columns to the correct fields, preview the data for errors, and "
    "import the records in bulk. This is especially useful for migrating historical data from "
    "other systems."
)

doc.add_heading("Additional Improvements", level=3)

improvements = [
    ("Dark mode — ", "An optional dark colour scheme that is easier on the eyes, toggled from the header."),
    ("Faster navigation — ", "Filtering and pagination on the submissions page will update instantly without reloading the entire page."),
    ("Password reset — ", "A \"Forgot Password\" option on the login page that sends a reset link via email."),
    ("Account security — ", "Automatic lockout after multiple failed login attempts, and configurable password expiry."),
]
for bold_part, rest in improvements:
    add_bullet(rest, bold_prefix=bold_part)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  5. IMPLEMENTATION TIMELINE
# ══════════════════════════════════════════════════════════

doc.add_heading("5. Implementation Timeline", level=1)

doc.add_paragraph(
    "The work is structured so that each phase delivers usable features upon completion. "
    "Phases are sequential, but the scope and priority of later phases can be adjusted based "
    "on feedback from earlier ones."
)

add_styled_table(
    ["Phase", "Focus Area", "Estimated Duration"],
    [
        ["Phase 1", "Online registration, approval workflow, user management, document uploads", "4 – 6 weeks"],
        ["Phase 2", "Enhanced profiles, certification tracking, voyage history, medical records", "3 – 4 weeks"],
        ["Phase 3", "Advanced search, audit log, email notifications, admin settings", "3 – 4 weeks"],
        ["Phase 4", "Reporting enhancements, data import, dark mode, password management", "2 – 3 weeks"],
        ["Total", "", "12 – 17 weeks"],
    ],
    col_widths=[1.0, 3.8, 1.7],
)

doc.add_paragraph(
    "Each phase will conclude with a demonstration and handover session. Any adjustments to scope "
    "or priority can be discussed between phases to ensure the system evolves in line with the "
    "Authority's operational needs."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  6. SUMMARY OF DELIVERABLES
# ══════════════════════════════════════════════════════════

doc.add_heading("6. Summary of Deliverables", level=1)

add_styled_table(
    ["#", "Deliverable", "Phase", "Priority"],
    [
        ["1", "Online Seafarer Registration Form", "1", "Critical"],
        ["2", "Application Review & Approval Workflow", "1", "Critical"],
        ["3", "User Management Interface", "1", "Critical"],
        ["4", "Document Upload & Storage", "1", "Critical"],
        ["5", "Enhanced Seafarer Profile Information", "2", "High"],
        ["6", "Certification Lifecycle Tracking", "2", "High"],
        ["7", "Voyage & Employment History", "2", "High"],
        ["8", "Medical Records Management", "2", "High"],
        ["9", "Advanced Search & Filtering", "3", "Medium"],
        ["10", "Activity & Audit Log", "3", "Medium"],
        ["11", "Email Notification System", "3", "Medium"],
        ["12", "Administration Settings Page", "3", "Medium"],
        ["13", "Enhanced Reporting & Batch Printing", "4", "Standard"],
        ["14", "Spreadsheet Data Import Tool", "4", "Standard"],
        ["15", "Interface Improvements (Dark Mode, Speed)", "4", "Standard"],
        ["16", "Password Reset & Account Security", "4", "Standard"],
    ],
    col_widths=[0.4, 3.3, 0.8, 1.0],
)

doc.add_paragraph("")

# ── Closing ─────────────────────────────────────────────

doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— End of Proposal —")
run.italic = True
run.font.color.rgb = GRAY

doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "This proposal is submitted for review and approval. Implementation will proceed "
    "upon confirmation of scope and phase prioritisation."
)
run.font.size = Pt(10)
run.font.color.rgb = GRAY

doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("NSPD Ghana Development Team  |  April 2026")
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = BLUE

# ── Save ────────────────────────────────────────────────

output_path = os.path.join(os.path.dirname(__file__), "NSPD_Ghana_V2_Proposal.docx")
doc.save(output_path)
print(f"Proposal saved to: {output_path}")
